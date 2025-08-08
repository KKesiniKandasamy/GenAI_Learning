# Utilise libraries/Data Setup
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import tensorflow as tf
import tensorflow.keras as keras
import sklearn as sk


# Load the dataset
file_path = "enhanced_diabetes_dataset.csv"
df = pd.read_csv(file_path)

from docx import Document


# Identify and handle outliers
def detect_outliers_iqr(data, feature):
    Q1 = data[feature].quantile(0.25)
    Q3 = data[feature].quantile(0.75)
    IQR = Q3 - Q1
    lower_bound = Q1 - 1.5 * IQR
    upper_bound = Q3 + 1.5 * IQR
    return data[(data[feature] < lower_bound) | (data[feature] > upper_bound)]

# Create a new Word document
doc = Document()
doc.add_heading('Outlier Detection Report', level=1)

outliers = {}
features = df.select_dtypes(include='number').columns.tolist()
# Detect and collect outliers for each feature
for feature in features:
    outliers[feature] = detect_outliers_iqr(df, feature)

# Write outlier results to the document
for feature, outlier_data in outliers.items():
    if not outlier_data.empty:
        doc.add_heading(f"Outliers detected in {feature}:", level=2)
        doc.add_paragraph(outlier_data.to_string())
    else:
        doc.add_heading(f"No outliers detected in {feature}.", level=2)

# Save the document before modifying the DataFrame
doc.save("outlier_report.docx")

# Handle outliers by removing them
for feature in features:
    outlier_indices = detect_outliers_iqr(df, feature).index
    df = df.drop(index=outlier_indices)

# Reset index after dropping outliers
df.reset_index(drop=True, inplace=True)

# Optional: display cleaned DataFrame
print("\nDataFrame after removing outliers:\n", df.head())